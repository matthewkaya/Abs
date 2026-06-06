# Copyright (c) 2026 Automatia BCN. All rights reserved.
# Licensed under the Business Source License 1.1.
# Production use requires a Commercial License - see LICENSE.
# Change Date: 2030-05-07 -> Apache License, Version 2.0

"""T-011 — RAG ingest pipeline v10: parsing, late chunking, contextual prefix."""

from __future__ import annotations

import hashlib
import logging
import os
import re
import tempfile
import unicodedata
import uuid
from dataclasses import dataclass


# Founder Tester Round 2 (BUG-6, infra fix) — Qdrant point IDs must be
# unsigned ints or RFC-4122 UUIDs. The previous chunk_id format
# (`<doc_id>-<seq:04d>`) failed Qdrant validation with "is not a valid
# point ID". We derive a deterministic UUID5 from doc_id+seq so reruns
# remain idempotent and `id` is still inspectable from a known doc.
_CHUNK_NAMESPACE = uuid.UUID("8a3b9f2c-1e4d-4f5a-9c7e-2b1d3e4f5a6b")


def _chunk_uuid(doc_id: str, seq: int) -> str:
    return str(uuid.uuid5(_CHUNK_NAMESPACE, f"{doc_id}/{seq}"))

logger = logging.getLogger(__name__)

__all__ = [
    "Chunk",
    "ParsedDocument",
    "estimate_token_count",
    "late_chunks",
    "parse_document",
    "parse_text",
]


@dataclass(slots=True)
class ParsedDocument:
    doc_id: str
    text: str
    mime_type: str
    metadata: dict[str, str]


@dataclass(slots=True)
class Chunk:
    chunk_id: str
    text: str
    raw_text: str
    doc_id: str
    seq: int
    char_start: int
    char_end: int
    metadata: dict[str, str]


def estimate_token_count(text: str) -> int:
    return max(1, len(text) // 4 + text.count(" ") // 4)


# Zero-width / BOM / word-joiner code points that PDF + DOCX extractors leak
# into the text and that pollute embeddings without adding meaning.
_ZERO_WIDTH = dict.fromkeys(
    map(ord, "​‌‍⁠﻿­"), None
)
# Latin typographic ligatures that PDF text layers emit as single glyphs; left
# as-is they fragment a word ("ﬁnance" ≠ "finance") and hurt recall.
_LIGATURES = {
    "ﬀ": "ff",
    "ﬁ": "fi",
    "ﬂ": "fl",
    "ﬃ": "ffi",
    "ﬄ": "ffl",
    "ﬅ": "st",
    "ﬆ": "st",
}
_MULTISPACE_RE = re.compile(r"[ \t ]{2,}")
_TRAILING_WS_RE = re.compile(r"[ \t ]+\n")
_MANY_NEWLINES_RE = re.compile(r"\n{3,}")


def _normalize(text: str) -> str:
    if text.startswith("﻿"):
        text = text[1:]
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _clean_text(text: str) -> str:
    """Normalize extracted text so the vector store sees clean, consistent
    content (founder feedback 2026-06-06: "metni chunk yapmadan önce text'e
    çevirip okunmayan karakterleri temizlemelisin; direk okuduğun gibi atarsan
    vector yanlış çalışır" + Turkish-character handling).

    Steps, order matters:
      1. Unicode NFC — composes Turkish letters (i̇/ş/ğ/ç/ö/ü) into single code
         points so "şirket" embeds identically however the source encoded it.
      2. Expand Latin ligatures (ﬁ→fi) that PDF layers emit as one glyph.
      3. Drop zero-width / soft-hyphen / BOM noise.
      4. Drop the U+FFFD replacement char (the tell-tale of a mis-decoded byte).
      5. Strip control chars except newline/tab.
      6. Collapse runs of spaces and 3+ blank lines.
    """
    if not text:
        return ""
    text = unicodedata.normalize("NFC", text)
    for glyph, repl in _LIGATURES.items():
        if glyph in text:
            text = text.replace(glyph, repl)
    text = text.translate(_ZERO_WIDTH)
    text = text.replace("�", "")
    text = "".join(
        ch for ch in text if ch in "\n\t" or unicodedata.category(ch)[0] != "C"
    )
    text = _TRAILING_WS_RE.sub("\n", text)
    text = _MULTISPACE_RE.sub(" ", text)
    text = _MANY_NEWLINES_RE.sub("\n\n", text)
    return text.strip()


def _doc_id(source: str) -> str:
    return hashlib.sha1(source.encode("utf-8", errors="replace")).hexdigest()[:16]


def parse_text(
    content: bytes | str,
    *,
    mime_type: str = "text/plain",
    filename: str | None = None,
) -> ParsedDocument:
    if isinstance(content, bytes):
        text = content.decode("utf-8", errors="replace")
        size = len(content)
    else:
        text = content
        size = len(content.encode("utf-8", errors="replace"))
    text = _clean_text(_normalize(text))
    return ParsedDocument(
        doc_id=_doc_id(filename or text),
        text=text,
        mime_type=mime_type,
        metadata={"filename": filename or "", "size": str(size)},
    )


_TEXT_MIMES = {
    "text/plain",
    "text/markdown",
    "application/json",
}
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_BINARY_PARSER_MIMES = {
    "application/pdf",
    _DOCX_MIME,
    _XLSX_MIME,
}


def _extract_binary_text(content: bytes, mime_type: str) -> str:
    """Extract plain text from a PDF/DOCX byte payload.

    Prefers `unstructured` when available (richer layout handling), but falls
    back to the lightweight, pure-Python `pypdf` / `python-docx` so the
    customer image stays small and disk-friendly (unstructured pulls GBs of
    ML layout models). Raises RuntimeError with an actionable message if no
    parser is available, and surfaces an empty-text hint for scanned PDFs.
    """
    # Preferred path: unstructured (only if the operator installed it).
    try:
        from unstructured.partition.auto import partition  # type: ignore

        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            elements = partition(filename=tmp_path)
            return "\n\n".join(
                el.text for el in elements if getattr(el, "text", None)
            )
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except ImportError:
        pass  # fall through to the lightweight parsers

    import io

    if mime_type == "application/pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("pypdf is required for PDF parsing") from exc
        try:
            reader = PdfReader(io.BytesIO(content))
            pages = [(page.extract_text() or "") for page in reader.pages]
        except Exception as exc:  # noqa: BLE001 — surface as clean 422, not 500
            raise RuntimeError(f"pdf_parse_failed: {exc}") from exc
        text = "\n\n".join(p for p in pages if p.strip())
        if not text.strip():
            raise RuntimeError(
                "pdf_no_extractable_text: the PDF appears to be scanned/image-only "
                "(no text layer). OCR is required to ingest it."
            )
        return text

    if mime_type == _DOCX_MIME:
        try:
            import docx  # type: ignore  # python-docx
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required for DOCX parsing") from exc
        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as exc:  # noqa: BLE001 — surface as clean 422, not 500
            raise RuntimeError(f"docx_parse_failed: {exc}") from exc
        return "\n\n".join(p.text for p in document.paragraphs if p.text.strip())

    if mime_type == _XLSX_MIME:
        try:
            import openpyxl  # type: ignore
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("openpyxl is required for XLSX parsing") from exc
        try:
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        except Exception as exc:  # noqa: BLE001 — surface as clean 422, not 500
            raise RuntimeError(f"xlsx_parse_failed: {exc}") from exc
        # Flatten each sheet to "Sheet: name" + tab-joined non-empty rows so a
        # spreadsheet of company records becomes searchable prose.
        parts: list[str] = []
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" \t ".join(cells))
            if rows:
                parts.append(f"[Sheet: {ws.title}]\n" + "\n".join(rows))
        try:
            wb.close()
        except Exception:
            pass
        text = "\n\n".join(parts)
        if not text.strip():
            raise RuntimeError("xlsx_no_extractable_text: the spreadsheet has no readable cells.")
        return text

    raise RuntimeError(f"no_parser_for_mime: {mime_type}")


def parse_document(
    content: bytes,
    *,
    mime_type: str,
    filename: str | None = None,
) -> ParsedDocument:
    mime_type = mime_type.lower()
    if mime_type in _TEXT_MIMES:
        return parse_text(content, mime_type=mime_type, filename=filename)

    if mime_type in _BINARY_PARSER_MIMES:
        joined = _extract_binary_text(content, mime_type)
        return parse_text(joined, mime_type=mime_type, filename=filename)

    logger.warning(
        "rag_parser_unknown_mime mime=%s falling_back=text/plain", mime_type
    )
    return parse_text(content, mime_type=mime_type, filename=filename)


_CHARS_PER_TOKEN = 4
_DELIMS = (".", "!", "?", "\n\n")


def late_chunks(
    doc: ParsedDocument,
    *,
    target_chars: int = 400,
    overlap_chars: int = 80,
    target_tokens: int | None = None,
    overlap_tokens: int | None = None,
    contextual_prefix: str | None = None,
) -> list[Chunk]:
    """Sentence-aware chunking. The primary unit is CHARACTERS (founder
    feedback 2026-06-06: chunks should cap around ~400 chars — small, precise
    chunks retrieve better than 2k-char blocks). ``target_tokens`` /
    ``overlap_tokens`` remain accepted for backward compatibility and, when
    given, override the char targets (1 token ≈ 4 chars)."""
    if target_tokens is not None:
        target_chars = target_tokens * _CHARS_PER_TOKEN
    if overlap_tokens is not None:
        overlap_chars = overlap_tokens * _CHARS_PER_TOKEN
    target_chars = max(1, target_chars)
    overlap_chars = max(0, min(overlap_chars, target_chars - 1))
    step = max(1, target_chars - overlap_chars)
    text = doc.text
    length = len(text)
    chunks: list[Chunk] = []
    seq = 0
    start = 0

    while start < length:
        end = min(start + target_chars, length)
        window = text[start:end]

        # Trim to nearest sentence boundary within the last 20% of the window,
        # but never shrink below half the target window.
        min_keep = max(1, target_chars // 2)
        search_start = int(0.8 * len(window))
        best_trim: int | None = None
        for delim in _DELIMS:
            idx = window.rfind(delim, search_start)
            if idx == -1:
                continue
            candidate = idx + len(delim)
            if candidate < min_keep:
                continue
            if best_trim is None or candidate > best_trim:
                best_trim = candidate
        if best_trim is not None and end < length:
            end = start + best_trim
            window = text[start:end]

        raw = window.strip()
        if not raw:
            start += step
            continue

        if contextual_prefix:
            final = f"{contextual_prefix.strip()}\n\n{raw}"
        else:
            final = raw

        meta = dict(doc.metadata)
        meta["seq"] = str(seq)
        meta["char_start"] = str(start)
        meta["char_end"] = str(end)
        chunks.append(
            Chunk(
                chunk_id=_chunk_uuid(doc.doc_id, seq),
                text=final,
                raw_text=raw,
                doc_id=doc.doc_id,
                seq=seq,
                char_start=start,
                char_end=end,
                metadata=meta,
            )
        )
        seq += 1
        if end >= length:
            break
        start += step

    if chunks and len(chunks) > 1:
        # Merge a too-small trailing chunk into its predecessor. Threshold is
        # 1/8 of the target window (char-based; matches the legacy token math
        # when target_tokens was supplied since target_chars = tokens * 4).
        min_chars = min(256, max(4, target_chars // 8))
        if len(chunks[-1].raw_text) < min_chars:
            tail = chunks.pop()
            prev = chunks[-1]
            merged_raw = f"{prev.raw_text} {tail.raw_text}"
            prev.raw_text = merged_raw
            prev.text = (
                f"{contextual_prefix.strip()}\n\n{merged_raw}"
                if contextual_prefix
                else merged_raw
            )
            prev.char_end = tail.char_end
            prev.metadata["char_end"] = str(tail.char_end)

    return chunks
