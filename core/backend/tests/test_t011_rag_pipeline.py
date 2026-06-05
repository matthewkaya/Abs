"""T-011 — RAG ingest + query pipeline tests."""

from __future__ import annotations

import base64
import hashlib
import secrets
from datetime import datetime, timezone
from unittest.mock import MagicMock

import pytest
from fastapi.testclient import TestClient
from sqlmodel import Session

from app.api.v1 import rag as rag_routes
from app.auth.oauth.models import OAuthClient
from app.db.session import get_engine
from app.main import app
from app.rag import pipeline_v10 as pipe
from types import SimpleNamespace


@pytest.fixture(autouse=True)
def _install_fake_cerbos_for_rag():
    """T-012 — RAG endpoints now run a Cerbos pre-check; default ALLOW for tests."""

    class _AllowingCerbos:
        def check_resources(self, *, principal, resources):  # noqa: ANN001
            entry = SimpleNamespace(is_allowed=lambda action: True)
            return SimpleNamespace(
                results=[entry], failed=lambda: False, status_code=200
            )

        def close(self) -> None:
            return None

    app.state.cerbos_client = _AllowingCerbos()
    yield
    if hasattr(app.state, "cerbos_client"):
        delattr(app.state, "cerbos_client")


def _challenge(verifier: str) -> str:
    return base64.urlsafe_b64encode(
        hashlib.sha256(verifier.encode("ascii")).digest()
    ).rstrip(b"=").decode("ascii")


def _seed_client(client_id: str) -> None:
    with Session(get_engine()) as db:
        db.add(
            OAuthClient(
                client_id=client_id,
                redirect_uris="https://app.local/callback",
                allowed_scopes="openid profile",
                is_confidential=False,
                created_at=datetime.now(timezone.utc).replace(tzinfo=None),
            )
        )
        db.commit()


def _issue_token(
    c: TestClient,
    *,
    client_id: str,
    user_subject: str,
    tenant_id: str,
    roles: list[str],
) -> str:
    verifier = "v" * 64
    auth = c.get(
        "/oauth/authorize",
        params={
            "response_type": "code",
            "client_id": client_id,
            "redirect_uri": "https://app.local/callback",
            "code_challenge": _challenge(verifier),
            "code_challenge_method": "S256",
            "scope": "rag:query",
            "user_subject": user_subject,
            "tenant_id": tenant_id,
            "roles": ",".join(roles),
        },
        follow_redirects=False,
    )
    assert auth.status_code == 302, auth.text
    code = auth.headers["location"].split("code=", 1)[1].split("&", 1)[0]
    tok = c.post(
        "/oauth/token",
        data={
            "grant_type": "authorization_code",
            "client_id": client_id,
            "code": code,
            "redirect_uri": "https://app.local/callback",
            "code_verifier": verifier,
        },
    )
    assert tok.status_code == 200, tok.text
    return tok.json()["access_token"]


# ----- pipeline_v10 ----------------------------------------------------


def test_parse_text_bytes_normalizes_crlf_and_bom() -> None:
    raw = b"\xef\xbb\xbfhello\r\nworld"
    doc = pipe.parse_text(raw, filename="t.txt")
    assert doc.text == "hello\nworld"
    assert doc.mime_type == "text/plain"
    assert doc.metadata["filename"] == "t.txt"
    assert doc.metadata["size"] == str(len(raw))
    assert len(doc.doc_id) == 16
    int(doc.doc_id, 16)


def test_parse_document_unknown_mime_falls_back_to_text(
    caplog: pytest.LogCaptureFixture,
) -> None:
    with caplog.at_level("WARNING"):
        doc = pipe.parse_document(b"hi", mime_type="application/x-zip", filename="z.bin")
    assert doc.mime_type == "application/x-zip"
    assert doc.text == "hi"
    assert any("unknown_mime" in r.getMessage() for r in caplog.records)


def test_parse_document_pdf_corrupt_raises_clean_runtimeerror() -> None:
    # A non-PDF / truncated payload must surface a clean RuntimeError (which the
    # /ingest-file route maps to 422), never a raw pypdf exception (→ 500).
    with pytest.raises(RuntimeError) as exc:
        pipe.parse_document(b"%PDF-1.4 not really a pdf", mime_type="application/pdf", filename="x.pdf")
    msg = str(exc.value).lower()
    assert "pdf_parse_failed" in msg or "pdf_no_extractable_text" in msg


def test_parse_document_docx_roundtrip() -> None:
    import io

    import docx  # python-docx

    buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Kira bedeli her ayin 5'inde odenir.")
    d.add_paragraph("Aidat 250 euro.")
    d.save(buf)
    doc = pipe.parse_document(
        buf.getvalue(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="lease.docx",
    )
    assert "Kira bedeli" in doc.text
    assert "Aidat 250" in doc.text


def test_parse_document_pdf_no_text_layer_raises() -> None:
    import io

    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    with pytest.raises(RuntimeError) as exc:
        pipe.parse_document(buf.getvalue(), mime_type="application/pdf", filename="scan.pdf")
    assert "pdf_no_extractable_text" in str(exc.value)


def test_late_chunks_basic_count_and_ordering() -> None:
    text = ". ".join(["sentence " * 20 for _ in range(40)])
    doc = pipe.parse_text(text, filename="big.txt")
    chunks = pipe.late_chunks(doc, target_tokens=200, overlap_tokens=20)
    assert len(chunks) >= 2
    assert [c.seq for c in chunks] == list(range(len(chunks)))
    seen_ids: set[str] = set()
    for c in chunks:
        # Founder Tester Round 2 (BUG-6 infra fix) — chunk_id is now a
        # deterministic UUID5 (Qdrant requires UUID/uint point IDs).
        # The ID must still be stable per (doc_id, seq) and unique
        # across the chunk list.
        assert c.chunk_id == pipe._chunk_uuid(doc.doc_id, c.seq)
        assert c.chunk_id not in seen_ids
        seen_ids.add(c.chunk_id)
        assert c.char_end > c.char_start


def test_late_chunks_overlap_present() -> None:
    text = ". ".join(["sentence " * 15 for _ in range(60)])
    doc = pipe.parse_text(text, filename="x.txt")
    chunks = pipe.late_chunks(doc, target_tokens=150, overlap_tokens=30)
    for prev, nxt in zip(chunks, chunks[1:]):
        assert nxt.char_start < prev.char_end


def test_late_chunks_contextual_prefix_applied() -> None:
    text = ". ".join(["sentence " * 8 for _ in range(80)])
    doc = pipe.parse_text(text, filename="ctx.txt")
    prefix = "DOC SUMMARY"
    chunks = pipe.late_chunks(
        doc, target_tokens=100, overlap_tokens=10, contextual_prefix=prefix
    )
    assert chunks
    for c in chunks:
        assert c.text.startswith(f"{prefix}\n\n")
        assert prefix not in c.raw_text


def test_late_chunks_short_text_single_chunk() -> None:
    doc = pipe.parse_text("a short note.", filename="s.txt")
    chunks = pipe.late_chunks(doc, target_tokens=200, overlap_tokens=20)
    assert len(chunks) == 1
    assert chunks[0].raw_text == doc.text


def test_late_chunks_last_chunk_reaches_end() -> None:
    text = ". ".join(["sentence " * 10 for _ in range(40)])
    doc = pipe.parse_text(text, filename="end.txt")
    chunks = pipe.late_chunks(doc, target_tokens=180, overlap_tokens=20)
    assert chunks[-1].char_end == len(doc.text)


def test_estimate_token_count_monotonic() -> None:
    s = pipe.estimate_token_count("a" * 100)
    l = pipe.estimate_token_count("a" * 1000)
    assert l > s > 0


# ----- /v1/rag routes --------------------------------------------------


def test_ingest_text_happy_path(monkeypatch: pytest.MonkeyPatch) -> None:
    cid = f"rag-{secrets.token_hex(3)}"
    _seed_client(cid)

    upserted: dict = {}

    def fake_upsert(*, collection, tenant_id, points):
        upserted["collection"] = collection
        upserted["tenant_id"] = tenant_id
        upserted["count"] = len(points)
        return len(points)

    monkeypatch.setattr(rag_routes.qc, "ensure_collection", lambda *a, **k: None)
    monkeypatch.setattr(rag_routes.qc, "upsert_points", fake_upsert)

    payload = {"text": "hello world. " * 200, "filename": "foo.txt"}
    with TestClient(app) as c:
        token = _issue_token(
            c,
            client_id=cid,
            user_subject="alice",
            tenant_id="tenant-1",
            roles=["member"],
        )
        r = c.post(
            "/v1/rag/ingest",
            json=payload,
            headers={"Authorization": f"Bearer {token}", "X-ABS-Audience": cid},
        )

    assert r.status_code == 200, r.text
    body = r.json()
    assert body["chunks"] >= 1
    assert len(body["doc_id"]) == 16
    assert upserted["tenant_id"] == "tenant-1"
    assert upserted["count"] == body["chunks"]


def test_ingest_file_docx_happy_path(monkeypatch: pytest.MonkeyPatch) -> None:
    import io

    import docx  # python-docx

    cid = f"rag-{secrets.token_hex(3)}"
    _seed_client(cid)

    upserted: dict = {}

    def fake_upsert(*, collection, tenant_id, points):
        upserted["tenant_id"] = tenant_id
        upserted["count"] = len(points)
        return len(points)

    monkeypatch.setattr(rag_routes.qc, "ensure_collection", lambda *a, **k: None)
    monkeypatch.setattr(rag_routes.qc, "upsert_points", fake_upsert)

    buf = io.BytesIO()
    d = docx.Document()
    for _ in range(30):
        d.add_paragraph("Kira bedeli her ayin 5'inde odenir. Aidat 250 euro. ")
    d.save(buf)

    with TestClient(app) as c:
        token = _issue_token(
            c, client_id=cid, user_subject="alice", tenant_id="tenant-1", roles=["member"]
        )
        # Browsers send octet-stream for .docx — the endpoint resolves the MIME
        # from the .docx extension, so this must still route to the DOCX parser.
        r = c.post(
            "/v1/rag/ingest-file",
            files={"file": ("lease.docx", buf.getvalue(), "application/octet-stream")},
            headers={"Authorization": f"Bearer {token}", "X-ABS-Audience": cid},
        )

    assert r.status_code == 200, r.text
    body = r.json()
    assert body["chunks"] >= 1
    assert upserted["tenant_id"] == "tenant-1"
    assert upserted["count"] == body["chunks"]


def test_ingest_file_empty_returns_400() -> None:
    cid = f"rag-{secrets.token_hex(3)}"
    _seed_client(cid)
    with TestClient(app) as c:
        token = _issue_token(
            c, client_id=cid, user_subject="alice", tenant_id="tenant-1", roles=["member"]
        )
        r = c.post(
            "/v1/rag/ingest-file",
            files={"file": ("empty.txt", b"", "text/plain")},
            headers={"Authorization": f"Bearer {token}", "X-ABS-Audience": cid},
        )
    assert r.status_code == 400


def test_ingest_file_scanned_pdf_returns_422() -> None:
    import io

    from pypdf import PdfWriter

    cid = f"rag-{secrets.token_hex(3)}"
    _seed_client(cid)
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    with TestClient(app) as c:
        token = _issue_token(
            c, client_id=cid, user_subject="alice", tenant_id="tenant-1", roles=["member"]
        )
        r = c.post(
            "/v1/rag/ingest-file",
            files={"file": ("scan.pdf", buf.getvalue(), "application/pdf")},
            headers={"Authorization": f"Bearer {token}", "X-ABS-Audience": cid},
        )
    # Scanned/no-text PDF → clean 422, not a 500.
    assert r.status_code == 422, r.text
    assert "pdf_no_extractable_text" in r.text


def test_ingest_text_missing_tenant_returns_403(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    cid = f"rag-{secrets.token_hex(3)}"
    _seed_client(cid)
    monkeypatch.setattr(rag_routes.qc, "ensure_collection", lambda *a, **k: None)
    monkeypatch.setattr(rag_routes.qc, "upsert_points", MagicMock())

    with TestClient(app) as c:
        token = _issue_token(
            c,
            client_id=cid,
            user_subject="bob",
            tenant_id="",
            roles=["member"],
        )
        r = c.post(
            "/v1/rag/ingest",
            json={"text": "data", "filename": "x.txt"},
            headers={"Authorization": f"Bearer {token}", "X-ABS-Audience": cid},
        )
    assert r.status_code == 403
    assert r.json()["detail"] == "missing_tenant_claim"


def test_query_returns_hits(monkeypatch: pytest.MonkeyPatch) -> None:
    cid = f"rag-{secrets.token_hex(3)}"
    _seed_client(cid)
    fake_hits = [
        {
            "id": "d-0001",
            "score": 0.92,
            "payload": {
                "chunk_id": "d-0001",
                "doc_id": "d",
                "seq": 0,
                "text": "hello",
                "tenant_id": "tenant-1",
            },
        }
    ]
    monkeypatch.setattr(rag_routes.qc, "ensure_collection", lambda *a, **k: None)
    monkeypatch.setattr(rag_routes.qc, "search", lambda *a, **k: fake_hits)

    with TestClient(app) as c:
        token = _issue_token(
            c,
            client_id=cid,
            user_subject="carol",
            tenant_id="tenant-1",
            roles=["member"],
        )
        r = c.post(
            "/v1/rag/query",
            json={"query": "hello", "limit": 3},
            headers={"Authorization": f"Bearer {token}", "X-ABS-Audience": cid},
        )
    assert r.status_code == 200, r.text
    body = r.json()
    assert len(body["hits"]) == 1
    assert body["hits"][0]["text"] == "hello"
    assert body["hits"][0]["score"] == 0.92
    assert isinstance(body["elapsed_ms"], (int, float))


def test_query_missing_jwt_returns_401() -> None:
    with TestClient(app) as c:
        r = c.post("/v1/rag/query", json={"query": "test", "limit": 5})
    assert r.status_code == 401
    assert r.json()["detail"] == "missing_bearer_token"


def test_ingest_empty_text_rejected_by_validation(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    cid = f"rag-{secrets.token_hex(3)}"
    _seed_client(cid)
    monkeypatch.setattr(rag_routes.qc, "ensure_collection", lambda *a, **k: None)
    monkeypatch.setattr(rag_routes.qc, "upsert_points", MagicMock())

    with TestClient(app) as c:
        token = _issue_token(
            c,
            client_id=cid,
            user_subject="dave",
            tenant_id="tenant-1",
            roles=["member"],
        )
        r = c.post(
            "/v1/rag/ingest",
            json={"text": "", "filename": "empty.txt"},
            headers={"Authorization": f"Bearer {token}", "X-ABS-Audience": cid},
        )
    assert r.status_code == 422
    errors = r.json()["detail"]
    assert any(err.get("loc")[-1] == "text" for err in errors)
